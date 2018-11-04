#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import RGBColor
import datetime
import xlsxwriter
import requests
from requests.auth import HTTPDigestAuth
import json
import getpass

requests.packages.urllib3.disable_warnings() 

def set_column_width(column, width):
    column.width = width
    for cell in column.cells:
        cell.width = width
     
########################					########################
########################	LTM Functions	########################
########################					########################

def create_excel_ltm(file_name, partitions, route_domain, device_details, provisioned_modules, self_ips, virtual_servers, pools, persistence, monitor, profile, ssl_cert, trunk, routes, vlans, rules):

	# Create a workbook and add a worksheet.
	workbook = xlsxwriter.Workbook(file_name)
	xls_device_details = workbook.add_worksheet('Device Details')
	xls_virtual = workbook.add_worksheet('vServers')
	xls_pool = workbook.add_worksheet('Pools')
	xls_profiles = workbook.add_worksheet('Profiles')
	xls_persistence = workbook.add_worksheet('Persistence')
	xls_monitors = workbook.add_worksheet('Monitors')
	xls_self_ips = workbook.add_worksheet('Self IPs')
	xls_ssl_certs = workbook.add_worksheet('Certs')
	xls_partitions = workbook.add_worksheet('Partitions')
	xls_rules = workbook.add_worksheet('iRules')



	cell_format = workbook.add_format()
	cell_format.set_text_wrap()
	xls_device_details.set_column('B:B', 40)
	xls_device_details.set_column('A:A', 15)
	xls_ssl_certs.set_column(0,4, 25)
	xls_virtual.set_column(0,4, 15)
	xls_pool.set_column(0,5, 15)
	xls_pool.set_column('E:E', 25)
	xls_self_ips.set_column(0,8, 25)
	xls_monitors.set_column(0,8, 25)
	xls_persistence.set_column(0,8, 25)
	xls_profiles.set_column(0,8, 25)
	xls_partitions.set_column(0,8, 20)
	xls_rules.set_column(0,8, 20)

	
	
	##########		Certificates  & SSL Profiles 		#################
	row = 0
	col = 0
	xls_ssl_certs.write(row, col,     'Certificates')
	xls_ssl_certs.write(row+1, col,     '--------------------------------------------')

	row += 1
	xls_ssl_certs.write(row, col,     'Name')
	xls_ssl_certs.write(row, col + 1,  'Partition')
	xls_ssl_certs.write(row, col + 2,  'Created By')
	xls_ssl_certs.write(row, col + 3,  'Expiration Time')
	xls_ssl_certs.write(row, col + 4,  'Create Time')

	row = 1
	 
	# Iterate over the data and write it out row by row.
	for key in ssl_cert:
		xls_ssl_certs.write(row, col,     key['name'])
		xls_ssl_certs.write(row, col + 1,  key['partition'])
		xls_ssl_certs.write(row, col + 2,  key['createdBy'])
		xls_ssl_certs.write(row, col + 3,  key['expirationString'])
		xls_ssl_certs.write(row, col + 4,  key['createTime'])
		row += 1


	row += 5
	col = 0
	xls_ssl_certs.write(row, col,     'SSL Profiles')
	xls_ssl_certs.write(row+1, col,     '--------------------------------------------')

	xls_ssl_certs.write(row, col,     'SSL Profiles Configuration')
	row += 1
	xls_ssl_certs.write(row, col,     'Name')
	xls_ssl_certs.write(row, col + 1,  'Partition')
	xls_ssl_certs.write(row, col + 2,  'Parent Profile')
	xls_ssl_certs.write(row, col + 3,  'Ciphers')
	xls_ssl_certs.write(row, col + 4,  'Cipher Group')
	row += 1


	for key in profile['ssl']:
		xls_ssl_certs.write(row, col, key['name'])
		xls_ssl_certs.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_ssl_certs.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,'\n'))
		xls_ssl_certs.write(row, col + 3, key['ciphers'])
		xls_ssl_certs.write(row, col + 4, key['cipherGroup'])
		row += 1




	# Write a total using a formula.


	xls_device_details.write(0, 0, 'Platform')
	xls_device_details.write(1, 0, 'Hostname')
	xls_device_details.write(2, 0, 'Version')
	xls_device_details.write(3, 0, 'Chassis ID')
	xls_device_details.write(4, 0, 'DNS Servers')
	xls_device_details.write(5, 0, 'NTP Servers')
	xls_device_details.write(6, 0, 'Syslog Servers')
	xls_device_details.write(7, 0, 'Time Zone')
	xls_device_details.write(10, 0, 'Active Modules')

	xls_device_details.write(11, 0, 'LTM')
	xls_device_details.write(12, 0, 'ASM')
	xls_device_details.write(13, 0, 'APM')
	xls_device_details.write(14, 0, 'DNS')
	xls_device_details.write(15, 0, 'AVR')
	xls_device_details.write(16, 0, 'AFM')
	xls_device_details.write(17, 0, 'Failover State')
	xls_device_details.write(18, 0, 'Failover Unicast IP')
	xls_device_details.write(19, 0, 'Management IP')
	xls_device_details.write(20, 0, 'ConfigSync IP')
	xls_device_details.write(21, 0, 'Mirror IP')
	xls_device_details.write(22, 0, 'Mirror IP (Secondary)')


	xls_device_details.write(0, 1, device_details['marketingName'])
	xls_device_details.write(1, 1, device_details['hostname'])
	xls_device_details.write(2, 1, device_details['version'])
	xls_device_details.write(3, 1, device_details['chassisId'])
	xls_device_details.write(4, 1, device_details['nameServers'])
	xls_device_details.write(5, 1, device_details['ntpServers'])
	xls_device_details.write(6, 1, device_details['remoteServers'])
	xls_device_details.write(7, 1, device_details['timeZone'])


	activeModules =""
	for key in device_details['activeModules']:
		activeModules += key.replace("|", "\n -")
		activeModules += "\n"

	xls_device_details.write(10, 1, activeModules, cell_format)
	xls_device_details.write(11, 1, provisioned_modules['ltm'])
	xls_device_details.write(12, 1, provisioned_modules['asm'])
	xls_device_details.write(13, 1, provisioned_modules['apm'])
	xls_device_details.write(14, 1, provisioned_modules['gtm'])
	xls_device_details.write(15, 1, provisioned_modules['avr'])
	xls_device_details.write(16, 1, provisioned_modules['afm'])


	xls_device_details.write(17, 1, device_details['failoverState'])
	xls_device_details.write(18, 1, device_details['unicastAddress'])
	xls_device_details.write(19, 1, device_details['managementIp'])
	xls_device_details.write(20, 1, device_details['configsyncIp'])
	xls_device_details.write(21, 1, device_details['mirrorIP'])
	xls_device_details.write(22, 1, device_details['mirrorSecondaryIp'])



	# Start from the first cell. Rows and columns are zero indexed.
	row = 0
	col = 0

	xls_self_ips.write(row, col,     'Self IP addresses')	
	row += 1
	xls_self_ips.write(row, col,     'Name')
	xls_self_ips.write(row, col + 1,  'Address')
	xls_self_ips.write(row, col + 2,  'Partition')
	xls_self_ips.write(row, col + 3,  'Vlan')
	xls_self_ips.write(row, col + 4,  'Floating')
	xls_self_ips.write(row, col + 5,  'Allow Service')


	row = 1

	# Iterate over the data and write it out row by row.
	for key in self_ips:
		xls_self_ips.write(row, col, key['name'])
		xls_self_ips.write(row, col + 1, key['address'])
		remove_partition = "/" + key['partition'] + "/"
		xls_self_ips.write(row, col + 2, key['partition'])
		xls_self_ips.write(row, col + 3, key['vlan'].replace(remove_partition,''))
		xls_self_ips.write(row, col + 4, key['floating'])
		if(type(key['allowService'])==list):
			allowService = ', '.join(key['allowService'])
		else:
			allowService = key['allowService']
		xls_self_ips.write(row, col + 5, allowService)
		row += 1

	# Write a total using a formula.



	#		Print Virtual Servers 			#row = 0
	col = 0
	row = 0
	xls_virtual.write(row, col,     'Virtual Servers Overview')
	xls_virtual.write(row+1, col,     '---------------------------------------')
	row += 1
	
	xls_virtual.write(row, col,     'Name')
	xls_virtual.write(row, col + 1,  'IP Address')
	xls_virtual.write(row, col + 2,  'Partition')

	row = 1


	for key in virtual_servers:
		xls_virtual.write(row, col, key['name'])
		remove_partition = "/" + key['partition'] + "/"
		xls_virtual.write(row, col + 1, key['destination'].replace(remove_partition,''))
		xls_virtual.write(row, col + 2, key['partition'])
		row += 1

	col = 0
	row += 5
	xls_virtual.write(row, col,     'Virtual Servers Details')
	xls_virtual.write(row+1, col,     '---------------------------------------')
	row += 1
	
	xls_virtual.write(row+1, col,  'Name')
	xls_virtual.write(row+2, col,  'Partition')
	xls_virtual.write(row+3, col,  'SubPath')
	xls_virtual.write(row+4, col,  'IP Address')
	xls_virtual.write(row+5, col,  'Mask')
	xls_virtual.write(row+6, col,  'Pool')
	xls_virtual.write(row+7, col,  'Source')
	xls_virtual.write(row+8, col,  'iRules Attached')
	xls_virtual.write(row+9, col,  'Policies Attached')
	xls_virtual.write(row+10, col,  'Profiles Attached')
	xls_virtual.write(row+11, col,  'Persistence')
	xls_virtual.write(row+12, col,  'Vlans status')
	xls_virtual.write(row+13, col,  'Vlans')
	xls_virtual.write(row+14, col,  'Syn cookies')
	xls_virtual.write(row+15, col,  'Enabled')
	xls_virtual.write(row+16, col,  'IP Protocol')
	xls_virtual.write(row+17, col,  'Rate limit')
	xls_virtual.write(row+18, col,  'Throughput Capacity (Mbps)')
	xls_virtual.write(row+19, col,  'Service Down Action')
	xls_virtual.write(row+20, col,  'Auto Last Hop')
	xls_virtual.write(row+21, col,  'Source Port')	
	xls_virtual.write(row+22, col,  'Translate Address')	
	xls_virtual.write(row+23, col,  'Translate Port')	
	xls_virtual.write(row+24, col,  'Eviction Policy')	
	
	
	col = 1

	for key in virtual_servers:

		xls_virtual.write(row+1, col,  key['name'])
		xls_virtual.write(row+2, col,  key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_virtual.write(row+3, col,  key['subPath'])
		if (len(key['subPath'])>0):
			 remove_path = "/" + key['partition'] + "/" + key['subPath'] + "/"
		else:
			 remove_path = "/" + key['partition'] + "/"
		xls_virtual.write(row+4, col,  key['destination'].replace(remove_partition,''))
		xls_virtual.write(row+5, col,  key['mask'])
		xls_virtual.write(row+6, col,  key['pool'].replace(remove_path,''))
		xls_virtual.write(row+7, col,  key['source'])
		if(type(key['rules'])==list):
			vs_rules = '\n'.join(key['rules']).replace(remove_partition,'')
		else:
			vs_rules = key['rules']
		xls_virtual.write(row+8, col,  vs_rules, cell_format)
		if(type(key['policyName'])==list):
			policyName = '\n'.join(key['policyName'])
		else:
			policyName = key['policyName']
		xls_virtual.write(row+9, col,  policyName, cell_format)
		if(type(key['profileName'])==list):
			profileName = '\n'.join(key['profileName'])
		else:
			profileName = key['profileName']	
		xls_virtual.write(row+10, col,  profileName, cell_format)
		if(type(key['securityLogProfiles'])==list):
			securityLogProfiles = '\n'.join(key['securityLogProfiles']).replace(remove_partition,'')
		else:
			securityLogProfiles = key['securityLogProfiles']	
		xls_virtual.write(row+11, col,  securityLogProfiles, cell_format)
		xls_virtual.write(row+12, col, key['persistence'])
		xls_virtual.write(row+13, col, key['vlanStatus'])
		xls_virtual.write(row+14, col, key['vlan'])
		xls_virtual.write(row+15, col, key['synCookieStatus'])
		xls_virtual.write(row+16, col, key['enabled'])
		xls_virtual.write(row+17, col, key['ipProtocol'])
		xls_virtual.write(row+18, col, key['rateLimit'])
		xls_virtual.write(row+19, col, key['throughputCapacity'])
		xls_virtual.write(row+20, col, key['serviceDownImmediateAction'])
		xls_virtual.write(row+21, col, key['autoLasthop'])
		xls_virtual.write(row+22, col, key['sourcePort'])
		xls_virtual.write(row+23, col, key['translateAddress'])
		xls_virtual.write(row+24, col, key['translatePort'])
		xls_virtual.write(row+25, col, key['flowEvictionPolicy'])
		col += 1

	################		Print Pools		################
	col = 0
	row = 0
	xls_pool.write(row, col,     'Pool Configuration')	
	row += 1

	xls_pool.write(row, col,     'Name')
	xls_pool.write(row, col + 1,  'Partition')
	xls_pool.write(row, col + 2,  'Members')
	xls_pool.write(row, col + 3,  'LB algorithm')
	xls_pool.write(row, col + 4,  'Monitors')
	row = 1


	for key in pools:
		xls_pool.write(row, col, key['name'])
		xls_pool.write(row, col + 1, key['partition'])
		if (len(key['subPath'])>0):
			 remove_path = "/" + key['partition'] + "/" + key['subPath'] + "/"
		else:
			 remove_path = "/" + key['partition'] + "/"
		if(type(key['members'])==list):
			members = '\n'.join(key['members'])
		else:
			members = key['members']	
		xls_pool.write(row, col + 2,  members, cell_format)
		xls_pool.write(row, col + 3, key['loadBalancingMode'])
		if 'min' in key['pool_monitor']:
			my_monitors = key['pool_monitor'].replace(remove_path,'\n')
		else:
			pool_monitor = key['pool_monitor']	
			my_monitors = pool_monitor.replace(remove_path,'')
		xls_pool.write(row, col + 4,  my_monitors, cell_format)		
			
		row += 1



	################		Print Source Peristence		################
	col = 0
	row = 0
	xls_persistence.write(row, col,     'Name')
	xls_persistence.write(row, col + 1,  'Partition')
	xls_persistence.write(row, col + 2,  'Parent Profile')
	xls_persistence.write(row, col + 3,  'Timeout')
	xls_persistence.write(row, col + 4,  'Mask')
	xls_persistence.write(row, col + 5,  'Mirroring')
	xls_persistence.write(row, col + 6,  'Match Across Pools')
	xls_persistence.write(row, col + 7,  'Match Across virtuals')
	row = 1


	for key in persistence['source']:
		xls_persistence.write(row, col, key['name'])
		xls_persistence.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_persistence.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,'\n'))
		xls_persistence.write(row, col + 3, key['timeout'])
		xls_persistence.write(row, col + 4, key['mask'])
		xls_persistence.write(row, col + 5, key['mirror'])
		xls_persistence.write(row, col + 6, key['matchAcrossPools'])
		xls_persistence.write(row, col + 7, key['matchAcrossVirtuals'])
		row += 1

		
	################	 Print Cookie Peristence		################

	col = 0
	row += 5
	xls_persistence.write(row, col,     'Name')
	xls_persistence.write(row, col + 1,  'Partition')
	xls_persistence.write(row, col + 2,  'Parent Profile')
	xls_persistence.write(row, col + 3,  'Encrypt Cookies')
	xls_persistence.write(row, col + 4,  'Encrypted Cookies')
	xls_persistence.write(row, col + 5,  'Encrypt Pool Name')
	row += 1


	for key in persistence['cookie']:
		xls_persistence.write(row, col, key['name'])
		xls_persistence.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_persistence.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,'\n'))
		xls_persistence.write(row, col + 3, key['cookieEncryption'])
		xls_persistence.write(row, col + 4, key['cookieName'])
		xls_persistence.write(row, col + 5, key['encryptCookiePoolname'])
		row += 1


		
	################	 Print HTTP Monitors		################
	row = 0
	col = 0

	xls_monitors.write(row, col,     'Name')
	xls_monitors.write(row, col + 1,  'Partition')
	xls_monitors.write(row, col + 2,  'Parent Monitor')
	xls_monitors.write(row, col + 3,  'Interval')
	xls_monitors.write(row, col + 4,  'Timeout')
	xls_monitors.write(row, col + 5,  'Send String')
	xls_monitors.write(row, col + 6,  'Receive String')
	row += 1


	for key in monitor['http']:
		xls_monitors.write(row, col, key['name'])
		xls_monitors.write(row, col + 1, key['partition'])
		xls_monitors.write(row, col + 2, key['defaultsFrom'])
		xls_monitors.write(row, col + 3, key['interval'])
		xls_monitors.write(row, col + 4, key['timeout'])
		xls_monitors.write(row, col + 5, key['send'])
		xls_monitors.write(row, col + 6, key['recv'])
		row += 1

	################	 Print HTTPS Monitors		################
	row += 5
	col = 0

	xls_monitors.write(row, col,     'Name')
	xls_monitors.write(row, col + 1,  'Partition')
	xls_monitors.write(row, col + 2,  'Parent Monitor')
	xls_monitors.write(row, col + 3,  'Interval')
	xls_monitors.write(row, col + 4,  'Timeout')
	xls_monitors.write(row, col + 5,  'Send String')
	xls_monitors.write(row, col + 6,  'Receive String')
	row += 1


	for key in monitor['https']:
		xls_monitors.write(row, col, key['name'])
		xls_monitors.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_monitors.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,'\n'))
		xls_monitors.write(row, col + 3, key['interval'])
		xls_monitors.write(row, col + 4, key['timeout'])
		xls_monitors.write(row, col + 5, key['send'])
		xls_monitors.write(row, col + 6, key['recv'])
		row += 1


	################	 Print Other Monitors		################
	row += 5
	col = 0

	xls_monitors.write(row, col,     'Name')
	xls_monitors.write(row, col + 1,  'Partition')
	xls_monitors.write(row, col + 2,  'Parent Monitor')
	xls_monitors.write(row, col + 3,  'Type')
	xls_monitors.write(row, col + 4,  'Interval')
	xls_monitors.write(row, col + 5,  'Timeout')
	row += 1


	for key in monitor['other']:
		xls_monitors.write(row, col, key['name'])
		xls_monitors.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_monitors.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,'\n'))
		xls_monitors.write(row, col + 3, key['proto'])
		xls_monitors.write(row, col + 4, key['interval'])
		xls_monitors.write(row, col + 5, key['timeout'])
		row += 1



	################	 Print HTTP Profiles 		################
	row = 0
	col = 0

	xls_profiles.write(row, col,     'HTTP Profiles Configuration')
	row += 1

	xls_profiles.write(row, col,     'Name')
	xls_profiles.write(row, col + 1,  'Partition')
	xls_profiles.write(row, col + 2,  'Parent Profile')
	xls_profiles.write(row, col + 3,  'Encrypt Cookies')
	xls_profiles.write(row, col + 4,  'Fallback Host')
	xls_profiles.write(row, col + 5,  'X-Forwarded-For')
	xls_profiles.write(row, col + 6,  'HSTS')
	row += 1



	for key in profile['http']:
		xls_profiles.write(row, col, key['name'])
		xls_profiles.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_profiles.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,''))
		if(len(key['encryptCookies'])>0):
			encryptCookies = '\n'.join(key['encryptCookies'])	
		else:
			encryptCookies = "none"	
		xls_profiles.write(row, col + 3, encryptCookies, cell_format)
		xls_profiles.write(row, col + 4, key['fallbackHost'])
		xls_profiles.write(row, col + 5, key['xff'])
		xls_profiles.write(row, col + 6, key['hsts'])
		row += 1

			
	################	 Print TCP Profiles 		################
	row += 5
	col = 0

	xls_profiles.write(row, col,     'TCP Profiles Configuration')
	row += 1
	xls_profiles.write(row, col,     'Name')
	xls_profiles.write(row, col + 1,  'Partition')
	xls_profiles.write(row, col + 2,  'Parent Profile')
	xls_profiles.write(row, col + 3,  'Idle Timeout')
	xls_profiles.write(row, col + 4,  'Keep Alive Interval')
	row += 1


	for key in profile['tcp']:
		xls_profiles.write(row, col, key['name'])
		xls_profiles.write(row, col + 1, key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_profiles.write(row, col + 2, key['defaultsFrom'].replace(remove_partition,'\n'))
		xls_profiles.write(row, col + 3, key['idleTimeout'])
		xls_profiles.write(row, col + 4, key['keepAliveInterval'])
		row += 1

		
	################		Print Route Domains	################
	row = 0
	col = 0
	xls_partitions.write(row, col,     'Route Domains')
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------------------------')
	row += 2
	
	xls_partitions.write(row, col,     'Name')
	xls_partitions.write(row, col + 1,  'Default Route Domain')


	row += 1
	 
	# Iterate over the data and write it out row by row.
	for key in partitions:
		xls_partitions.write(row, col,     key['name'])
		xls_partitions.write(row, col + 1,  key['defaultRouteDomain'])
		row += 1

	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------------------------')
	################		Print Partitions		################
	
	row += 4
	col = 0
	xls_partitions.write(row, col,     'Partitions')
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------')
	row += 2
	
	xls_partitions.write(row, col,     'Name')
	xls_partitions.write(row, col + 1,  'Partition')
	xls_partitions.write(row, col + 2,  'Parent')
	xls_partitions.write(row, col + 3,  'ID')
	xls_partitions.write(row, col + 4,  'VLAN')
	xls_partitions.write(row, col + 5,  'Eviction Policy')
	xls_partitions.write(row, col + 6,  'Connection Limit')
	xls_partitions.write(row, col + 7,  'Throughput Limit')

	row += 1
	 
	# Iterate over the data and write it out row by row.
	for key in route_domain:
		xls_partitions.write(row, col,     key['name'])
		xls_partitions.write(row, col + 1,  key['partition'])
		remove_partition = "/" + key['partition'] + "/"
		xls_partitions.write(row, col + 2,  key['parent'])
		xls_partitions.write(row, col + 3,  key['id'])
		if(type(key['vlans'])==list):
			rd_vlans = '\n'.join(key['vlans']).replace(remove_partition,'')
		else:
			rd_vlans = key['vlans']	
		xls_partitions.write(row, col +4,  rd_vlans, cell_format)		
		xls_partitions.write(row, col + 5,  key['flowEvictionPolicy'].replace(remove_partition,''))
		xls_partitions.write(row, col + 6,  key['connectionLimit'])
		xls_partitions.write(row, col + 7,  key['throughputCapacity'])
		row += 1
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------------------------')		
			
	################		Print Trunks		################
		
	row += 4
	xls_partitions.write(row, col,     'Trunks')
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------')
	row += 2
	
	xls_partitions.write(row, col,     'Name')
	xls_partitions.write(row, col + 1,  'Distribution Hash')
	xls_partitions.write(row, col + 2,  'LACP')
	xls_partitions.write(row, col + 3,  'Media')
	xls_partitions.write(row, col + 4,  'Interfaces')

	row += 1

	# Iterate over the data and write it out row by row.
	for key in trunk:

		if(type(key['interfaces'])==list):
			trunk_interfaces = '\n'.join(key['interfaces'])
		else:
			trunk_interfaces = key['interfaces']	
		
		xls_partitions.write(row, col,     key['name'])
		xls_partitions.write(row, col + 1,  key['distributionHash'])
		xls_partitions.write(row, col + 2,  key['lacp'])
		xls_partitions.write(row, col + 3,  key['media'])
		xls_partitions.write(row, col + 4,  trunk_interfaces)
		row += 1
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------------------------')		
		
	################		Print Routes		################
		
	row += 4
	xls_partitions.write(row, col,     'Routes')
	xls_partitions.write(row+1, col,     '---------------------------------------------------------------------------------------------------')
	row += 2
	
	xls_partitions.write(row, col,     'Name')
	xls_partitions.write(row, col + 1,  'Partition')
	xls_partitions.write(row, col + 2,  'Network')
	xls_partitions.write(row, col + 3,  'Type')
	xls_partitions.write(row, col + 4,  'Resource')
	xls_partitions.write(row, col + 5,  'MTU')

	row += 1
	 
	# Iterate over the data and write it out row by row.
	for key in routes:
		xls_partitions.write(row, col,     key['name'])
		xls_partitions.write(row, col + 1,  key['partition'])
		xls_partitions.write(row, col + 2,  key['network'])
		xls_partitions.write(row, col + 3,  key['type'])
		xls_partitions.write(row, col + 4,  key['resource'])
		xls_partitions.write(row, col + 5,  key['mtu'])
		row += 1
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------------------------')		
		
	################		Print VLANs		################
		
	row += 4
	xls_partitions.write(row, col,     'VLANs')
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------')
	row += 2
	
	xls_partitions.write(row, col,     'Name')
	xls_partitions.write(row, col + 1,  'Partition')
	xls_partitions.write(row, col + 2,  'Auto Last Hop')
	xls_partitions.write(row, col + 3,  'Failsafe')
	xls_partitions.write(row, col + 4,  'Failsafe Action')
	xls_partitions.write(row, col + 5,  'Failsafe Timeout')
	xls_partitions.write(row, col + 6,  'MTU')
	xls_partitions.write(row, col + 7,  'Interfaces')

	row += 1
	 
	# Iterate over the data and write it out row by row.
	for key in vlans:
		xls_partitions.write(row, col, key['name'])
		xls_partitions.write(row, col + 1,  key['partition'])
		xls_partitions.write(row, col + 2,  key['autoLasthop'])
		xls_partitions.write(row, col + 3,  key['failsafe'])
		xls_partitions.write(row, col + 4,  key['failsafeAction'])
		xls_partitions.write(row, col + 5,  key['failsafeTimeout'])
		xls_partitions.write(row, col + 6,  key['mtu'])
		xls_partitions.write(row, col + 7,  key['interfaces'])
		
		row += 1
	xls_partitions.write(row+1, col,     '------------------------------------------------------------------------------------------------------')		
		
	################	 Print iRules		################
	row = 0
	col = 0

	xls_rules.write(row, col,     'iRules')
	xls_rules.write(row+1, col,     '------------------------------------------------------------------------------------------------------')		
	row += 2

	xls_rules.write(row, col, "iRule Name")
	xls_rules.write(row, col+1, "Partition")
	xls_rules.write(row, col+2, "iRule")	

	for key in rules:
		xls_rules.write(row, col, key['name'])
		xls_rules.write(row, col+1, key['partition'])
		xls_rules.write(row, col+2, key['apiAnonymous'], cell_format)
		row += 1


	workbook.close()

	
def word_ltm_overview(document, device_details, suggestions):

	document.add_heading('F5 LTM Configuration Review', level=1)

	document.add_paragraph('During the meetings with "customer" and the configuration reviews we did on the F5 appliances we present our findings in this report.')

	document.add_heading('Overview', level=2)
	document.add_paragraph('The following table provides athe summary of all the recommendations of the health check for this ASM policy.')
	
	num_of_suggestions = 0
	table = document.add_table(rows=1, cols=4)

	table.style = 'Table Grid'
	table.cell(0,0).text = '#'
	table.cell(0,1).text = 'Severity'
	table.cell(0,2).text = 'Suggestion'
	table.cell(0,3).text = 'Category'
	
	for key in suggestions:
		if key['severity'] == "error":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']

	for key in suggestions:
		if key['severity'] == "warning":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']
	
	for key in suggestions:
		if key['severity'] == "info":
			num_of_suggestions +=1
			cells = table.add_row().cells
			cells[0].text = str(num_of_suggestions)
			paragraph = cells[1].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			cells[2].text = key['txt']
			cells[3].text = key['section']
	
	
	set_column_width(table.columns[0], Cm(1))		
	set_column_width(table.columns[1], Cm(1))		
	set_column_width(table.columns[2], Cm(12))		
	set_column_width(table.columns[3], Cm(3))		

	document.add_paragraph()
	document.add_paragraph()
	
	document.add_heading('Device Details', level=2)
	document.add_paragraph('The following table shows the general configuration from the devices that have been configured.')


	activeModules =""
	for key in device_details['activeModules']:
		activeModules += key.replace("|", "\n -")
		activeModules += "\n"
		
	table = document.add_table(rows=10, cols=2)
	table.style = 'Table Grid'
	table.cell(0,0).text = 'Details'
	table.cell(1,0).text = 'Platform'
	table.cell(2,0).text = 'Hostname'
	table.cell(3,0).text = 'Version'
	table.cell(4,0).text = 'Chassis ID'
	table.cell(5,0).text = 'DNS Servers'
	table.cell(6,0).text = 'NTP Servers'
	table.cell(7,0).text = 'Syslog Servers'
	table.cell(8,0).text = 'Time Zone'
	table.cell(9,0).text = 'Licensed Modules'
#------------------------------------------------------------------
	table.cell(0,1).text = 'Values'
	table.cell(1,1).text = device_details['marketingName']
	table.cell(2,1).text = device_details['hostname']
	table.cell(3,1).text = device_details['version']
	table.cell(4,1).text = device_details['chassisId']
	table.cell(5,1).text = device_details['nameServers']
	table.cell(6,1).text = device_details['ntpServers']
	table.cell(7,1).text = device_details['remoteServers']
	table.cell(8,1).text = device_details['timeZone']
	table.cell(9,1).text = activeModules
	

	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Overview":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_modules(document, provisioned_modules):

	document.add_heading('Provisioned Modules', level=2)
	document.add_paragraph('The following table shows the general configuration from the devices that have been configured.')


	table = document.add_table(rows=7, cols=2)
	table.style = 'Table Grid'
	table.cell(0,0).text = 'Modules'
	table.cell(1,0).text = 'Local Traffic (LTM)'
	table.cell(2,0).text = 'Application Security (ASM)'
	table.cell(3,0).text = 'Access Policy (APM)'
	table.cell(4,0).text = 'Global Traffic (DNS)'
	table.cell(5,0).text = 'Application Visibility and Reporting (AVR)'
	table.cell(6,0).text = 'Advanced Firewall (AFM)'
#------------------------------------------------------------------
	table.cell(0,1).text = 'Values'
	table.cell(1,1).text = provisioned_modules['ltm']
	table.cell(2,1).text = provisioned_modules['asm']
	table.cell(3,1).text = provisioned_modules['apm']
	table.cell(4,1).text = provisioned_modules['gtm']
	table.cell(5,1).text = provisioned_modules['avr']
	table.cell(6,1).text = provisioned_modules['afm']


	document.add_paragraph()
	document.add_paragraph()


	document.save("reports/F5 LTM - Config Review.docx")	
def word_ltm_ha(document, device_details, suggestions):

	document.add_heading('HA Configuration', level=2)
	document.add_paragraph('The following table shows the general configuration from the devices that have been configured.')


	table = document.add_table(rows=7, cols=2)
	table.style = 'Table Grid'
	#table.columns[0].width = Inches(1)
	table.cell(0,0).text = 'HA Configuration'
	table.cell(1,0).text = 'Failover State'
	table.cell(2,0).text = 'Failover Unicast IP'
	table.cell(3,0).text = 'Management IP'
	table.cell(4,0).text = 'ConfigSync IP '
	table.cell(5,0).text = 'Mirror IP'
	table.cell(6,0).text = 'Mirror IP (Secondary)'
#------------------------------------------------------------------
	table.cell(0,1).text = 'Value'
	table.cell(1,1).text = device_details['failoverState']
	table.cell(2,1).text = device_details['unicastAddress']
	table.cell(3,1).text = device_details['managementIp']
	table.cell(4,1).text = device_details['configsyncIp']
	table.cell(5,1).text = device_details['mirrorIP']
	table.cell(6,1).text = device_details['mirrorSecondaryIp']


	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="HA Configuration":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
	

	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_self_ips(document, self_ips, suggestions):

	document.add_heading('SelfIPs', level=2)
	document.add_paragraph('A virtual server is one of the most important components of any BIG-IP Local Traffic Manager (LTM) configuration. A virtual server type allows for a full-proxy service—that is, a client-side connection and a server-side conn.')

	table = document.add_table(rows=1, cols=6)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'SelfIP Name'
	table.cell(0,1).text = 'IP Address'
	table.cell(0,2).text = 'Partition'
	table.cell(0,3).text = 'VLAN'
	table.cell(0,4).text = 'Floating'
	table.cell(0,5).text = 'Port Lockdown'


	for key in self_ips:
#		print (key)
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['address']
		cells[2].text = key['partition']
		remove_partition = "/" + key['partition'] + "/"
		cells[3].text = key['vlan'].replace(remove_partition,'')
		cells[4].text = key['floating']
		if(type(key['allowService'])==list):
			cells[5].text = ', '.join(key['allowService'])
		else:
			cells[5].text = key['allowService']		
		
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="SelfIPs":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
	

	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_virtual_servers(document, virtual_servers, suggestions):

	document.add_heading('Virtual Servers', level=2)
	document.add_paragraph('A virtual server is one of the most important components of any BIG-IP Local Traffic Manager (LTM) configuration. A virtual server type allows for a full-proxy service—that is, a client-side connection and a server-side connection with data passing between them. F5 supports multiple types of virtual server types but recommends Standard virtual servers whenever L7 intelligence is or may required. During our audit we will be analyzing only Standard Virtual Servers. For more information on the other virtual server types, refer to AskF5 article (https://support.f5.com/csp/article/K14163).')

	table = document.add_table(rows=1, cols=3)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Virtual Server Name'
	table.cell(0,1).text = 'Address:Port'
	table.cell(0,2).text = 'Partition'


	for key in virtual_servers:
#		print (key)
		cells = table.add_row().cells
		if (len(key['subPath'])>0):
			 remove_path = "/" + key['partition'] + "/" + key['subPath'] + "/"
		else:
			 remove_path = "/" + key['partition'] + "/"		
		cells[0].text = key['name']
		cells[1].text = key['destination'].replace(remove_path,'')
		cells[2].text = key['partition']

		
	document.add_paragraph()
	document.add_paragraph('In the following sections we provide more details on the configuration of each of the above Virtual Servers')

	for key in virtual_servers:
		document.add_heading(key['name'], level=3)

		table = document.add_table(rows=24, cols=2)
		table.style = 'Table_rows'
		cells = table.add_row().cells
		table.cell(0,0).text = 'Configuration'
		table.cell(1,0).text = 'Virtual Server Name'
		table.cell(2,0).text = 'IP Address:Port'
		table.cell(3,0).text = 'Partition'
		table.cell(4,0).text = 'Path'
		table.cell(5,0).text = 'Mask'
		table.cell(6,0).text = 'Pool'
		table.cell(7,0).text = 'iRules Attached'
		table.cell(8,0).text = 'Policies Attached'
		table.cell(9,0).text = 'Profiles Attached'
		table.cell(10,0).text = 'Log Profiles Attached'
		table.cell(11,0).text = 'Persistence'
		table.cell(12,0).text = 'Vlans status'
		table.cell(13,0).text = 'Vlans'
		table.cell(14,0).text = 'Syn cookies'	
		table.cell(15,0).text = 'Enabled'	
		table.cell(16,0).text = 'IP Protocol'	
		table.cell(17,0).text = 'Rate limit'	
		table.cell(18,0).text = 'Throughput Capacity (Mbps)'	
		table.cell(19,0).text = 'Service Down Action'	
		table.cell(20,0).text = 'Auto Last Hop'	
		table.cell(21,0).text = 'Source Port'	
		table.cell(22,0).text = 'Translate Address'	
		table.cell(23,0).text = 'Translate Port'	
		table.cell(24,0).text = 'Eviction Policy'
		
		if (len(key['subPath'])>0):
			 remove_path = "/" + key['partition'] + "/" + key['subPath'] + "/"
			 remove_partition = "/" + key['partition'] + "/"
		else:
			 remove_path = "/" + key['partition'] + "/"	
			 remove_partition = "/" + key['partition'] + "/"
		table.cell(0,1).text = 'Value'
		table.cell(1,1).text = key['name']
		table.cell(2,1).text = key['destination'].replace(remove_partition,'')
		table.cell(3,1).text = key['partition']
		table.cell(4,1).text = key['subPath']
		table.cell(5,1).text = key['mask']
		table.cell(6,1).text = key['pool'].replace(remove_path,'')
		if(type(key['rules'])==list):
			table.cell(7,1).text = '\n'.join(key['rules']).replace(remove_partition,'')
		else:
			table.cell(7,1).text = key['rules'].replace(remove_partition,'')
		if(type(key['policyName'])==list):
			table.cell(8,1).text = '\n'.join(key['policyName'])
		else:
			table.cell(8,1).text = key['policyName']	
		if(type(key['profileName'])==list):
			table.cell(9,1).text = '\n'.join(key['profileName'])
		else:
			table.cell(9,1).text = key['profileName']			
		if(type(key['securityLogProfiles'])==list):
			table.cell(10,1).text = '\n'.join(key['securityLogProfiles']).replace(remove_partition,'')
		else:
			table.cell(10,1).text = key['securityLogProfiles']	
		table.cell(11,1).text = key['persistence']
		table.cell(12,1).text = key['vlanStatus']
		table.cell(13,1).text = key['vlan']
		table.cell(14,1).text = key['synCookieStatus']	
		table.cell(15,1).text = key['enabled']	
		table.cell(16,1).text = key['ipProtocol']	
		table.cell(17,1).text = key['rateLimit']	
		table.cell(18,1).text = str(key['throughputCapacity'])	
		table.cell(19,1).text = key['serviceDownImmediateAction']	
		table.cell(20,1).text = key['autoLasthop']	
		table.cell(21,1).text = key['sourcePort']	
		table.cell(22,1).text = key['translateAddress']	
		table.cell(23,1).text = key['translatePort']	
		table.cell(24,1).text = key['flowEvictionPolicy']	
		
		document.add_paragraph()
		document.add_paragraph()
	
	
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Virtual Servers":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()



	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_pools(document, pools, suggestions):

	document.add_heading('Pools', level=2)
	document.add_paragraph('Pool is a configuration object that groups backend services (or else called “pool members”) together to receive and process network traffic in a fashion determined by a specified load balancing algorithm.')

	document.add_paragraph('The following table lists the configured pools on the BIG-IP device.')

	table = document.add_table(rows=1, cols=5)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Pool Name'
	table.cell(0,1).text = 'Partition'
	table.cell(0,2).text = 'Pool Member Names'
	table.cell(0,3).text = 'Load Balancing Algorithm'
	table.cell(0,4).text = 'Health Monitors'

	for key in pools:
#		print (key)
		cells = table.add_row().cells
		cells[0].text = key['name']
		if (len(key['subPath'])>0):
			 remove_path = "/" + key['partition'] + "/" + key['subPath'] + "/"
		else:
			 remove_path = "/" + key['partition'] + "/"		
		cells[1].text = key['partition']
		if(type(key['members'])==list):
			cells[2].text = '\n'.join(key['members'])
		else:
			cells[2].text = key['members']			
		cells[3].text = key['loadBalancingMode']
		if 'min' in key['pool_monitor']:
			cells[4].text = key['pool_monitor'].replace(remove_path,'\n')
		else:
			pool_monitor = key['pool_monitor']	
			cells[4].text = pool_monitor.replace(remove_path,'')
			
		
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Pools":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()



	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_persistence(document, persistence, suggestions):

	document.add_heading('Persistence Profiles', level=2)
	document.add_paragraph('Many applications served by BIG-IP LTM are session-based and require the client to be load balanced to the same pool member for the duration of that session. BIG-IP LTM can accomplish this requirement through persistence. When a client connects to a virtual server for the first time, a load balancing decision is made, then the configured persistence record is created for that client. All subsequent connections that the client makes to that virtual server are sent to the same pool member for the life of that persistence record.')

	document.add_heading('Source IP Persistence Profiles', level=3)
	document.add_paragraph('The following table lists the configured Source IP persistence profiles on the BIG-IP device.')

	table = document.add_table(rows=1, cols=5)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Partition'
	table.cell(0,2).text = 'Timeout'
	table.cell(0,3).text = 'Mask'
	table.cell(0,4).text = 'Mirroring'


	for key in persistence['source']:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['partition']
		cells[2].text = key['timeout']
		cells[3].text = key['mask']
		cells[4].text = key['mirror']

		
	document.add_paragraph()
	document.add_paragraph()

	
	
	

	document.add_heading('Cookie Persistence', level=3)
	document.add_paragraph('The following table lists the configured Cookie persistence profiles on the BIG-IP device.')

	table = document.add_table(rows=1, cols=5)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Partition'
	table.cell(0,2).text = 'Cookie Name'
	table.cell(0,3).text = 'Encrypt Cookie'
	table.cell(0,4).text = 'Encrypt Pool Name'


	for key in persistence['cookie']:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['partition']
		cells[2].text = key['cookieName']
		cells[3].text = key['cookieEncryption']
		cells[4].text = key['encryptCookiePoolname']

		
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Persistence":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1
			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()


	
	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_monitors(document, monitor, suggestions):

	document.add_heading('Monitors', level=2)
	document.add_paragraph('The BIG-IP system uses monitors to check whether or not pool members are eligible to service application traffic. Monitors periodically send specific requests to pool members and evaluate their health based on the members’ response or lack thereof. Monitors can make explicit requests to an application, causing the application to perform an action which, in turn, tests vital server resources of that application.')


	document.add_heading('HTTP Monitors', level=3)
	document.add_paragraph('The following table lists the configured HTTP monitors on the BIG-IP device. ')

	table = document.add_table(rows=1, cols=6)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Parent'
	table.cell(0,2).text = 'Interval'
	table.cell(0,3).text = 'Timeout'
	table.cell(0,4).text = 'Send String'
	table.cell(0,5).text = 'Receive String'


	for key in monitor['http']:
#		print(key)
		cells = table.add_row().cells
		cells[0].text = key['name']
		remove_partition = "/" + key['partition'] + "/"
		cells[1].text = key['defaultsFrom'].replace(remove_partition,'')
		cells[2].text =  str(key['interval'])
		cells[3].text =  str(key['timeout'])
		cells[4].text =key['send']
		cells[5].text = key['recv']
		
	document.add_paragraph()


	document.add_heading('HTTPS Monitors', level=3)
	document.add_paragraph('The following table lists the configured HTTPS monitors on the BIG-IP device. ')

	table = document.add_table(rows=1, cols=6)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Parent'	
	table.cell(0,2).text = 'Interval'
	table.cell(0,3).text = 'Timeout'
	table.cell(0,4).text = 'Send String'
	table.cell(0,5).text = 'Receive String'


	for key in monitor['https']:
		cells = table.add_row().cells
		cells[0].text = key['name']
		remove_partition = "/" + key['partition'] + "/"
		cells[1].text = key['defaultsFrom'].replace(remove_partition,'')
		cells[2].text = str(key['interval'])
		cells[3].text = str(key['timeout'])
		cells[4].text = key['send']
		cells[5].text = key['recv']
		
	document.add_paragraph()


	document.add_heading('TCP/ICMP Monitors', level=5)
	document.add_paragraph('The following table lists the configured ICMP/TCP monitors on the BIG-IP device. ')

	table = document.add_table(rows=1, cols=5)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Parent'	
	table.cell(0,2).text = 'Type'
	table.cell(0,3).text = 'Interval'
	table.cell(0,4).text = 'Timeout'


	for key in monitor['other']:
		cells = table.add_row().cells
		cells[0].text = key['name']
		remove_partition = "/" + key['partition'] + "/"
		cells[1].text = key['defaultsFrom'].replace(remove_partition,'')
		cells[2].text = key['proto']
		cells[3].text =  str(key['interval'])
		cells[4].text =  str(key['timeout'])
		
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Monitors":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()



	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_profiles(document, profile, suggestions):
	document.add_heading('Profiles', level=2)
	document.add_paragraph('Profile is an object that contains settings with values, for controlling the behavior of a particular type of network traffic, such as HTTP connections. Profiles also provide a way for you to enable connection and session persistence, and to manage client application authentication.')


	document.add_heading('HTTP Profiles', level=3)
	document.add_paragraph('The following table lists the configured HTTP profiles on the BIG-IP device. ')

	table = document.add_table(rows=1, cols=6)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Parent Profile'
	table.cell(0,2).text = 'Encrypt Cookies'
	table.cell(0,3).text = 'Fallback Host'
	table.cell(0,4).text = 'X-Forwarded-For'
	table.cell(0,5).text = 'HSTS'


	for key in profile['http']:
#		print(key)
		cells = table.add_row().cells
		cells[0].text = key['name']
		remove_partition = "/" + key['partition'] + "/"
		cells[1].text = key['defaultsFrom'].replace(remove_partition,'')
		if(len(key['encryptCookies'])>0):
			cells[2].text = '\n'.join(key['encryptCookies'])	
		else:
			cells[2].text = "none"	
		cells[3].text = key['fallbackHost']
		cells[4].text = key['xff']
		cells[5].text =	key['hsts']

	document.add_paragraph()
	
	

	document.add_heading('TCP Profiles', level=3)
	document.add_paragraph('The following table lists the configured TCP profiles on the BIG-IP device. ')

	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Parent Profile'
	table.cell(0,2).text = 'Idle Timeout'
	table.cell(0,3).text = 'Keep Alive Interval'



	for key in profile['tcp']:
#		print(key)
		cells = table.add_row().cells
		cells[0].text = key['name']
		remove_partition = "/" + key['partition'] + "/"
		cells[1].text = key['defaultsFrom'].replace(remove_partition,'')
		cells[2].text = str(key['idleTimeout'])
		cells[3].text = str(key['keepAliveInterval'])

	document.add_paragraph()
	
	
	document.add_heading('SSL Client Profiles', level=3)
	document.add_paragraph('The following table lists the configured SSL Client profiles on the BIG-IP device. ')

	table = document.add_table(rows=1, cols=4)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Parent Profile'
	table.cell(0,2).text = 'Ciphers'
	table.cell(0,3).text = 'Cipher Group'


	for key in profile['ssl']:
		cells = table.add_row().cells
		cells[0].text = key['name']
		remove_partition = "/" + key['partition'] + "/"
		cells[1].text = key['defaultsFrom'].replace(remove_partition,'')
		cells[2].text = key['ciphers']
		cells[3].text = key['cipherGroup']

	document.add_paragraph()
	
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Profiles":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()

	
	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_certs(document, ssl_cert, suggestions):
	document.add_heading('SSL Certificates', level=2)
	document.add_paragraph("Before you can configure an SSL profile, you must install one or more SSL certificates on the BIG-IP system. The SSL certificate can be either a self-signed certificate or a trusted Certificate Authority (CA) certificate. ")
	document.add_paragraph("A self-signed SSL certificate is a certificate that is signed by its own private key. BIG-IP software includes a self-signed SSL certificate named default, which the SSL profile can use to terminate SSL traffic. You can also use the Configuration utility pages to create or renew additional self-signed certificates. A CA certificate is an SSL certificate that is signed by a CA's private key. Using a CA certificate allows you to replace the self-signed certificate on each BIG-IP system with a trusted CA certificate, which is a certificate signed by a third party.")

	document.add_paragraph('The following table lists the installed SSL certificates on the BIG-IP device.')

	table = document.add_table(rows=1, cols=3)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Created By'
	table.cell(0,2).text = 'Expiration Time'


	for key in ssl_cert:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['createdBy']
		cells[2].text = key['expirationString']

		
	document.add_paragraph()

	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Certs":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
	
	
	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_partitions(document, partitions, suggestions):
	document.add_heading('Partitions', level=2)
	document.add_paragraph("The following sections shows the Partitions that have been configured for the BIG-IP device. ")

	table = document.add_table(rows=1, cols=2)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Default Route Domain'


	for key in partitions:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['defaultRouteDomain']

	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Partitions":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
	
	
	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_route_domains(document, route_domain, suggestions):
	document.add_heading('Route Domains', level=2)
	document.add_paragraph("The following sections shows the Route Domains that have been configured for the BIG-IP device. ")

	table = document.add_table(rows=1, cols=7)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Partition'
	table.cell(0,2).text = 'ID'
	table.cell(0,3).text = 'VLANs'
	table.cell(0,4).text = 'Strict'
	table.cell(0,5).text = 'Parent'
	table.cell(0,6).text = 'Eviction Limit'


	for key in route_domain:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['partition']
		cells[2].text = str(key['id'])
		remove_partition = "/" + key['partition'] + "/"
		if(type(key['vlans'])==list):
			cells[3].text = '\n'.join(key['vlans']).replace(remove_partition,'')
		else:
			cells[3].text = key['vlans']
		cells[4].text = key['strict']
		cells[5].text = key['parent']
		cells[6].text = key['flowEvictionPolicy'].replace(remove_partition,'')

	
			
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Route Domains":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()
	
	
	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_vlans(document, vlans, suggestions):
	document.add_heading('VLANs', level=2)
	document.add_paragraph("The following sections shows the VLANs that have been configured for the BIG-IP device. ")

	table = document.add_table(rows=1, cols=6)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Partition'
	table.cell(0,2).text = 'AutoLastHop'
	table.cell(0,3).text = 'Failsafe'
	table.cell(0,4).text = 'FailsafeAction'
	table.cell(0,5).text = 'Interfaces'

	for key in vlans:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['partition']
		cells[2].text = key['autoLasthop']
		cells[3].text = key['failsafe']
		cells[4].text = key['failsafeAction']
		cells[5].text = key['interfaces']

		
	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="VLANs":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_routes(document, routes, suggestions):
	document.add_heading('Routing', level=2)
	document.add_paragraph("The following sections shows the Routes that have been configured for the BIG-IP device. ")

	table = document.add_table(rows=1, cols=5)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Partition'
	table.cell(0,2).text = 'Network'
	table.cell(0,3).text = 'Type'
	table.cell(0,4).text = 'Resource'

	for key in routes:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['partition']
		cells[2].text = key['network']
		cells[3].text = key['type']
		cells[4].text = key['resource']

	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Routes":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_trunk(document, trunk, suggestions):
	document.add_heading('Trunks', level=2)
	document.add_paragraph("The following sections shows the Trunk that have been configured for the BIG-IP device. ")

	table = document.add_table(rows=1, cols=5)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'Name'
	table.cell(0,1).text = 'Distribution Hash'
	table.cell(0,2).text = 'LACP'
	table.cell(0,3).text = 'Media'
	table.cell(0,4).text = 'interfaces'

	for key in trunk:
		cells = table.add_row().cells
		cells[0].text = key['name']
		cells[1].text = key['distributionHash']
		cells[2].text = key['lacp']
		cells[3].text = key['media']
		cells[4].text = key['interfaces']

	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="Trunks":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()

	
	document.save("reports/F5 LTM - Config Review.docx")
def word_ltm_irules(document, rules, suggestions):
	document.add_heading('iRules', level=2)
	document.add_paragraph("The following sections shows the iRules that have been configured for the BIG-IP device. ")

	table = document.add_table(rows=1, cols=2)
	table.style = 'Table_rows'
	table.cell(0,0).text = 'iRule Name'
	table.cell(0,1).text = 'iRule Details'


	for key in rules:
		if "_sys_" not in key['name']:
			cells = table.add_row().cells
			cells[0].text = key['name']
			cells[1].text = key['apiAnonymous']

	document.add_paragraph()
	document.add_paragraph()

	num_of_suggestions = 0

	for key in suggestions:
		if key['section']=="iRules":
			p=document.add_paragraph(style="Intense Quote")
			r = p.add_run()
			r.add_picture(key['severity']+'.png',width=Inches(.25), height=Inches(.25))
			r.add_text(key['txt'])
			num_of_suggestions +=1

			
	if num_of_suggestions==0:
		p=document.add_paragraph(style="Intense Quote")
		r = p.add_run()
		r.add_picture('low.png',width=Inches(.25), height=Inches(.25))
		r.add_text(' No suggestions found.')
		
	document.add_paragraph()
	document.add_paragraph()

	document.save("reports/F5 LTM - Config Review.docx")
		